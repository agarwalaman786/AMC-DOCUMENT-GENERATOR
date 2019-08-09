import babel.numbers
import datetime
from tkinter import * #To import the GUI framework
import tkinter.messagebox #To performs sometypes of popups
import datetime #To import the datetime module
import sqlite3  #Library to manage the database
from functools import partial  #TO pass the arguments in the function by the command attribute
from docx import * #To perform the operation on the document
from tkcalendar import Calendar, DateEntry  #To diplay the calendar
from dateutil.relativedelta import relativedelta #To perform operation on the date
from tkinter import filedialog  #to save the file
from datetime import date
from docx.shared import Pt
from dateutil.relativedelta import relativedelta
from PIL import Image
import subprocess
import os

mypath = os.getcwd()
myfilepath = os.getcwd()
myfile2path = os.getcwd()
mylist = mypath.split('\\')
mylist.append('AMC_MANAGEMENT_TABLE1.db')
db_file = '\\'.join(mylist)

myfilelist = myfilepath.split('\\')
myfilelist.append('SAMPLEBILLSES1.docx')
myses1path = '\\'.join(myfilelist)
print(myses1path)
myfile1list = myfile2path.split('\\')
myfile1list.append('SAMPLEBILLSES2.docx')
myses2path = '\\'.join(myfile1list)
print(myses2path)

#To format the given ammount
def special_format(n):
    s, *d = str(n).partition(".")
    r = ",".join([s[x-2:x] for x in range(-3, -len(s), -2)][::-1] + [s[-3:]])
    return "".join([r] + d)


#Code to check whether the table exist or not
def checkTableExists(dbcon, tablename):
    dbcur = dbcon.cursor()
    dbcur.execute("""
    SELECT count(*) FROM sqlite_master WHERE type = 'table' AND name = '{0}'
        """.format(tablename.replace('\'', '\'\'')))
    if dbcur.fetchone()[0] == 1:
        dbcur.close()
        return True
    dbcur.close()
    return False

def validate_Date(date1,date2,date3,date4,date5,date6):
    myvar = True
    if date1 =="" or date1 =="None" or date2 =="" or date2 =="None" or date3 =="" or date3 =="None" or date4 =="" or date4 =="None" or date5 =="" or date5 =="None":
        tkinter.messagebox.showinfo("Empty","Date can't be empty")
        myvar =  False
        return myvar
    temp1 = date1.split('/')
    lengthOfTemp1 = len(temp1)
    temp2 = date2.split('/')
    lengthOfTemp2 = len(temp2)
    temp3 = date3.split('/')
    lengthOfTemp3 = len(temp3)
    temp4 = date4.split('/')
    lengthOfTemp4 = len(temp4)
    temp5 = date5.split('/')
    lengthOfTemp5 = len(temp5)
    temp6 = date6.split('/')
    lengthOfTemp6 = len(temp6)
    
    if lengthOfTemp1!= 3 or lengthOfTemp2!= 3  or lengthOfTemp3!= 3 or lengthOfTemp4!= 3 or lengthOfTemp5!= 3 or lengthOfTemp6!= 3:
        tkinter.messagebox.showinfo("Format","Date format should be dd/mm/yyyy")
        myvar = False
        return myvar
    if ((int(temp1[0])>=1 and int(temp1[0])<=31) and (int(temp1[1])>=1 and int(temp1[1])<=12))==False or ((int(temp2[0])>=1 and int(temp2[0])<=31) and (int(temp2[1])>=1 and int(temp2[1])<=12))==False or ((int(temp3[0])>=1 and int(temp3[0])<=31) and (int(temp3[1])>=1 and int(temp3[1])<=12))==False or ((int(temp4[0])>=1 and int(temp4[0])<=31) and (int(temp4[1])>=1 and int(temp4[1])<=12))==False or ((int(temp5[0])>=1 and int(temp5[0])<=31) and (int(temp5[1])>=1 and int(temp5[1])<=12))==False or ((int(temp6[0])>=1 and int(temp6[0])<=31) and (int(temp6[1])>=1 and int(temp6[1])<=12))==False:
        tkinter.messagebox.showinfo("Correct","Date should be correct")
        myvar = False
        return myvar
    return myvar

def validate_Date1(date1,date2,date3,date4,date5,date6,date7,date8):
    myvar = True
    if date1 =="" or date1 =="None" or date2 =="" or date2 =="None" or date3 =="" or date3 =="None" or date4 =="" or date4 =="None" or date5 =="" or date5 =="None" or date6 =="" or date6=="None" or date7 =="" or date7=="None" or date8 =="" or date8=="None":
        tkinter.messagebox.showinfo("Empty","Date can't be empty")
        myvar = False
        return myvar
    temp1 = date1.split('/')
    lengthOfTemp1 = len(temp1)
    temp2 = date2.split('/')
    lengthOfTemp2 = len(temp2)
    temp3 = date3.split('/')
    lengthOfTemp3 = len(temp3)
    temp4 = date4.split('/')
    lengthOfTemp4 = len(temp4)
    temp5 = date5.split('/')
    lengthOfTemp5 = len(temp5)
    temp6 = date6.split('/')
    lengthOfTemp6 = len(temp6)
    temp7 = date7.split('/')
    lengthOfTemp7 = len(temp7)
    temp8 = date8.split('/')
    lengthOfTemp8 = len(temp8)
    
    if lengthOfTemp1!= 3 or lengthOfTemp2!= 3  or lengthOfTemp3!= 3 or lengthOfTemp4!= 3 or lengthOfTemp5!= 3 or lengthOfTemp6!= 3 or lengthOfTemp7!= 3 or lengthOfTemp8!= 3:
        tkinter.messagebox.showinfo("Format","Date format should be dd/mm/yyyy")
        myvar = False
        return myvar
    if ((int(temp1[0])>=1 and int(temp1[0])<=31) and (int(temp1[1])>=1 and int(temp1[1])<=12))==False or ((int(temp2[0])>=1 and int(temp2[0])<=31) and (int(temp2[1])>=1 and int(temp2[1])<=12))==False or ((int(temp3[0])>=1 and int(temp3[0])<=31) and (int(temp3[1])>=1 and int(temp3[1])<=12))==False or ((int(temp4[0])>=1 and int(temp4[0])<=31) and (int(temp4[1])>=1 and int(temp4[1])<=12))==False or ((int(temp5[0])>=1 and int(temp5[0])<=31) and (int(temp5[1])>=1 and int(temp5[1])<=12))==False or ((int(temp6[0])>=1 and int(temp6[0])<=31) and (int(temp6[1])>=1 and int(temp6[1])<=12))==False or ((int(temp7[0])>=1 and int(temp7[0])<=31) and (int(temp7[1])>=1 and int(temp7[1])<=12))==False or ((int(temp8[0])>=1 and int(temp8[0])<=31) and (int(temp8[1])>=1 and int(temp8[1])<=12))==False:
        tkinter.messagebox.showinfo("Correct","Date should be correct")
        myvar = False
        return myvar
    return myvar

#This is the starting of the main window of the GUI    

def Bill_generator_function():
    Bill = Toplevel()
    #Bill.geometry("600x350+300+300")
    Bill.state('zoomed')
    
    #Function to check null values in the field in the bill
    def check_bill(bill, ammount1, ammount2, ponumber, SES_Number1, SES_Number2,Senior_Officer, Senior_Designation):
        if bill == "" or bill=="None" or ammount1 == "" or ammount1 == "None"  or ponumber == "" or ponumber=="None" or Senior_Officer =="" or Senior_Officer=="None" or Senior_Designation == "" or Senior_Designation =="None":
            tkinter.messagebox.showinfo("Field","Fields can't be empty")
            return False
        elif SES_Number1 == "" or SES_Number1=="None" and SES_Number2!="":
            tkinter.messagebox.showinfo("SES","SES Number 1 is mandatory")
            return False
        elif ammount1 == "" or ammount1 == "None" and ammount2!="":
            tkinter.messagebox.showinfo("Ammount","Ammount1 is mandatory")
            return False    
        
        elif SES_Number1!= "" and SES_Number1!= "None" and SES_Number2!= "" and SES_Number1!= "None" and ammount2 == "None" and ammount2=="":
            tkinter.messagebox.showinfo('Ammount',"Ammount2 can't be empty")
            return False
        else:
            #tkinter.messagebox.showinfo('Correct','All fields are ok')
            return True
        
    #Function to print Bills of AMC
    def Printing_Bills():
        confirm = tkinter.messagebox.askquestion('Warning',"Are You Sure???", icon='warning')
        if confirm == 'yes':
            Invoice_date = Entry_date1.get()
            Service_s_date = Entry_service_start_date1.get()
            
            Service_e_date = Entry_service_end_date1.get()
            SES1_start_date = Entry_SES1_start_date1.get()
            SES1_end_date = Entry_SES1_end_date1.get()
            SES2_start_date = Entry_SES2_start_date1.get()
            SES2_end_date = Entry_SES2_end_date1.get()
            PODATE = Entry_PO_Date1.get()
            Key_for_AMC = list_var.get()
            
            if(check_bill(Bill_var.get(), Ammount_Var1.get(), Ammount_Var2.get(), PO_Number_Var.get(), SES_Number_Var1.get(), SES_Number_Var2.get(),Senior_Officer.get(),Senior_Designation.get())):
                #print(SES_Number_Var2.get())
                if (SES_Number_Var1.get()!=None or SES_Number_Var1.get()!= "") and (SES_Number_Var2.get()=="" or SES_Number_Var2.get()=="None"):
                    if (validate_Date(PODATE,Invoice_date,Service_s_date,Service_e_date,SES1_start_date,SES1_end_date)):
                        
                        conn = sqlite3.connect(db_file)
                        if checkTableExists(conn, 'bill'):
                            pass
                        else :
                            conn.execute('create table bill (id INTEGER PRIMARY KEY autoincrement, PO_Number varchar(100),PO_DATE varchar(20), Invoice_Number varchar(100), Invoice_date varchar(20), Service_start_date varchar(20), Service_end_date varchar(20),SES_number1 varchar(60), SES1_Ammount bigint, SES1_start_date varchar(20),SES1_end_date varchar(20),SES_number2 varchar(60), SES2_Ammount bigint, SES2_start_date varchar(20),SES2_end_date varchar(20),Senior_officer varchar(100), Designation varchar(100))')   
                        conn.execute("INSERT INTO bill (PO_Number,PO_DATE,Invoice_Number, Invoice_date, Service_start_date, Service_end_date ,SES_number1 , SES1_Ammount, SES1_start_date ,SES1_end_date ,SES_number2, SES2_Ammount , SES2_start_date,SES2_end_date,Senior_officer, Designation)                        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(PO_Number_Var.get(),PODATE,Bill_var.get(),Invoice_date,Service_s_date,Service_e_date,SES_Number_Var1.get(),Ammount_Var1.get(),SES1_start_date, SES1_end_date,SES_Number_Var2.get(), Ammount_Var2.get(), SES2_start_date, SES2_end_date, Senior_Officer.get(), Senior_Designation.get()));
                        conn.commit()
                        #match1 = conn.execute("select *from AMC_MANAGE_TABLE where HARDWARE_NAME  = ?",(Key_for_AMC,))
                        c = conn.cursor()
                        c.execute("select *from AMC_MANAGE_TABLE where HARDWARE_NAME  = ?",(Key_for_AMC,))
                        match = c.fetchone()
                        empty_list = []
                        if match == None:
                            tkinter.messagebox.showinfo("Warning","Hardware name can't be none")
                        else:
                            for row in match:
                                empty_list.append(row)
                            today = date.today().strftime("%d/%m/%Y")
                            for i in range(1):
                                if SES_Number_Var2.get()=="" or SES_Number_Var2.get()=="None" and SES_Number_Var1.get()!="" or SES_Number_Var1.get()!="None":
                                    document = Document(myses1path)
                                    for paragraph in document.paragraphs:
                                        if 'TODAYDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('TODAYDATE',str(today))
                                        if 'HARDWARENAME' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('HARDWARENAME',str(empty_list[1]))
                                        if 'SESTARTDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SESTARTDATE',str(Service_s_date))
                                        if 'SEENDDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SEENDDATE',str(Service_e_date))
                                        if 'CONTRACTNUMBER' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('CONTRACTNUMBER',str(empty_list[0]))
                                        if 'CONTRACTSTARTDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('CONTRACTSTARTDATE',str(empty_list[2]))
                                        if 'CONTRACTENDDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('CONTRACTENDDATE',str(empty_list[3]))
                                        if 'VENDORCODE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('VENDORCODE',str(empty_list[4]))
                                        if 'BILLNUMBERX' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('BILLNUMBERX',str(Bill_var.get()))
                                        if 'BILLDATEX' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('BILLDATEX',str(Invoice_date))
                                        if 'SESNUMBER12' in paragraph.text:
                    
                                            paragraph.text = paragraph.text.replace('SESNUMBER12',str(SES_Number_Var1.get()))
                                        if 'AMMOUNT1SES' in paragraph.text:
                                            myammount = str(Ammount_Var1.get()).replace(",","")
                                            paragraph.text = paragraph.text.replace('AMMOUNT1SES',special_format(float(myammount)))
                                        if 'HHHHHHHHHHHH' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('HHHHHHHHHHHH',str(Senior_Officer.get()))
                                        if 'DDDDDDDDDDD' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('DDDDDDDDDDD',str(Senior_Designation.get()))
                                        if 'VENDORNAMEX' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('VENDORNAMEX',str(empty_list[5]))
                                        if 'SESSPEIOD' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SESSPEIOD',str(SES1_start_date))
                                        if 'SESEPEIOD' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SESEPEIOD',str(SES1_end_date))
                                        if 'PONUMBER' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('PONUMBER',str(PO_Number_Var.get()))
                                        if 'POSTARTDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('POSTARTDATE',str(PODATE))
                                        style = document.styles['Normal']
                                        font = style.font
                                        font.name = 'Calibri'
                                        font.size = Pt(10)
                                        paragraph.style = document.styles['Normal']
                                
                                    for table in document.tables:
                                        for row in table.rows:
                                            for cell in row.cells:
                                                for paragraph in cell.paragraphs:
                                                    if 'TODAYDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('TODAYDATE',str(today))
                                                    if 'HARDWARENAME' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('HARDWARENAME',str(empty_list[1]))
                                                    if 'SESTARTDATE1' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESTARTDATE1',"")
                                                        paragraph.add_run(str(Service_s_date)).font.size = Pt(7)
                                                    if 'SEENDDATE1' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SEENDDATE1',"")
                                                        paragraph.add_run(str(Service_e_date)).font.size = Pt(7)
                                                    if 'SESTARTDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESTARTDATE',str(Service_s_date))
                                                    #paragraph.add_run(str(Service_s_date)).font.size = Pt(7)
                                                    if 'SEENDDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SEENDDATE',str(Service_e_date))
                                                    #paragraph.add_run(str(Service_e_date)).font.size = Pt(7)
                                                    if 'CONTRACTNUMBER' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('CONTRACTNUMBER',str(empty_list[0]))
                                                    if 'CONTRACTSTARTDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('CONTRACTSTARTDATE',str(empty_list[2]))
                                                    if 'CONTRACTENDDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('CONTRACTENDDATE',str(empty_list[3]))
                                                    if 'VENDORCODE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('VENDORCODE',str(empty_list[4]))
                                                    if 'BILLNUMBERX' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('BILLNUMBERX',str(Bill_var.get()))
                                                    if 'BILLDATEX' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('BILLDATEX',str(Invoice_date))
                                                    if 'SESNUMBER12' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESNUMBER12',"")
                                                        paragraph.add_run(str(SES_Number_Var1.get())).font.size = Pt(9)
                                                    if 'AMMOUNT1SES' in paragraph.text:
                                                        #paragraph.text = paragraph.text.replace('AMMOUNT1SES',str(Ammount_Var1.get()))
                                                        paragraph.text = paragraph.text.replace('AMMOUNT1SES',"")
                                                        paragraph.add_run(special_format(float(Ammount_Var1.get()))).font.size = Pt(10)
                                                    if 'HHHHHHHHHHHH' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('HHHHHHHHHHHH',str(Senior_Officer.get()))
                                                    if 'DDDDDDDDDDDD' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('DDDDDDDDDDDD',str(Senior_Designation.get()))
                                                    if 'VENDORNAMEX' in paragraph.text:
                                                        #paragraph.text = paragraph.text.replace('VENDORNAMEX',str(i[5]))
                                                        paragraph.text = paragraph.text.replace('VENDORNAMEX',"")
                                                        paragraph.add_run(str(empty_list[5])).font.size = Pt(10)
                                                    if 'SESSPEIODXY' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESSPEIODXY-SESEPEIODXY',"")
                                                        paragraph.add_run(str(SES1_start_date) + ' - ' +str(SES1_end_date)).font.size = Pt(7.5)
                                                    if 'SESTXYARTDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESTXYARTDATE-SEEXYNDDATE',"")
                                                        paragraph.add_run(str(Service_s_date) + ' - ' + str(Service_e_date)).font.size = Pt(7.5)                                        
                                                    if 'SESSPEIOD' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESSPEIOD',str(SES1_start_date))
                                                        #paragraph.text = paragraph.text.replace('SESSPEIOD',"")
                                                        #paragraph.add_run(str(SES1_start_date)).font.size = Pt(8)
                                                    if 'SESEPEIOD' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESEPEIOD',str(SES1_end_date))
                                                          #paragraph.text = paragraph.text.replace('SESEPEIOD',"")
                                                        #paragraph.add_run(str(SES1_end_date)).font.size = Pt(8)
                                                    if 'PONUMBER' in paragraph.text:
                                                        #paragraph.text = paragraph.text.replace('PONUMBER',str(PO_Number_Var.get()))
                                                        paragraph.text = paragraph.text.replace('PONUMBER',"")
                                                        paragraph.add_run(str(PO_Number_Var.get())).font.size = Pt(7.5)
                                                    if 'POSTARTDATE' in paragraph.text:
                                                        #paragraph.text = paragraph.text.replace('POSTARTDATE',str(PODATE))
                                                        paragraph.text = paragraph.text.replace('POSTARTDATE',"")
                                                        paragraph.add_run(str(PODATE)).font.size = Pt(8)
                                         
                                    file = filedialog.asksaveasfile(title = "Save a bill as", filetypes = (("word file","*.docx"),("word file","*.doc")))
                                    new_file = file.name
                                    document.save(new_file)
                                    conn.close()
                                    mylist1 = str(new_file).split('/')
                                    name = mylist1.pop()
                                    mylist = str(new_file).split('/')
                                    mylist.pop()
                                    myjoinlist = '/'.join(mylist)
                                    os.chdir(str(myjoinlist))
                                    subprocess.Popen('start winword ' + str(name),shell = True)
                                    confirm1 = tkinter.messagebox.askquestion('Warning',"Do you want to close the window???", icon='warning')
                                    if confirm1 == 'yes':
                                        Bill.after(1000, lambda: Bill.destroy())
                        
                        #This is for when the document is containing two SES Numbers
                else:
                    if (validate_Date1(PODATE,Invoice_date,Service_s_date,Service_e_date,SES1_start_date,SES1_end_date,SES2_start_date,SES2_end_date) ):
                        conn = sqlite3.connect(db_file)
                        if checkTableExists(conn, 'bill'):
                            pass
                        else :
                            conn.execute('create table bill (id INTEGER PRIMARY KEY autoincrement, PO_Number varchar(100),PO_DATE varchar(20), Invoice_Number varchar(100), Invoice_date varchar(20), Service_start_date varchar(20), Service_end_date varchar(20),SES_number1 varchar(60), SES1_Ammount bigint, SES1_start_date varchar(20),SES1_end_date varchar(20),SES_number2 varchar(60), SES2_Ammount bigint, SES2_start_date varchar(20),SES2_end_date varchar(20),Senior_officer varchar(100), Designation varchar(100))')   
                        conn.execute("INSERT INTO bill (PO_Number,PO_DATE,Invoice_Number, Invoice_date, Service_start_date, Service_end_date ,SES_number1 , SES1_Ammount, SES1_start_date ,SES1_end_date ,SES_number2, SES2_Ammount , SES2_start_date,SES2_end_date,Senior_officer, Designation)                        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)",(PO_Number_Var.get(),PODATE,Bill_var.get(),Invoice_date,Service_s_date,Service_e_date,SES_Number_Var1.get(),Ammount_Var1.get(),SES1_start_date, SES1_end_date,SES_Number_Var2.get(), Ammount_Var2.get(), SES2_start_date, SES2_end_date, Senior_Officer.get(), Senior_Designation.get()));
                        conn.commit()
                        c = conn.cursor()
                        c.execute("select *from AMC_MANAGE_TABLE where HARDWARE_NAME  = ?",(Key_for_AMC,))
                        match = c.fetchone()
                        empty_list = []
                        if match == None:
                            tkinter.messagebox.showinfo("Warning","Hardware name can't be none")
                        else:
                            for row in match:
                                empty_list.append(row)
                            
                            today = date.today().strftime("%d/%m/%Y")
                            for i in range(1):
                                if SES_Number_Var2.get()!="" or SES_Number_Var2.get()!="None" and SES_Number_Var1.get()!="" or SES_Number_Var1.get()!="None":
                                    document = Document(myses2path)
                                    for paragraph in document.paragraphs:
                                        if 'TODAYDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('TODAYDATE',str(today))
                                        if 'HARDWARENAME' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('HARDWARENAME',str(empty_list[1]))
                                        if 'SESTARTDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SESTARTDATE',str(Service_s_date))    
                                        if 'SEENDDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SEENDDATE',str(Service_e_date))
                                        if 'CONTRACTNUMBER' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('CONTRACTNUMBER',str(empty_list[0]))
                                        if 'CONTRACTSTARTDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('CONTRACTSTARTDATE',str(empty_list[2]))
                                        if 'CONTRACTENDDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('CONTRACTENDDATE',str(empty_list[3]))
                                        if 'VENDORCODE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('VENDORCODE',str(empty_list[4]))
                                        if 'BILLNUMBERX' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('BILLNUMBERX',str(Bill_var.get()))
                                        if 'BILLDATEX' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('BILLDATEX',str(Invoice_date))
                                        if 'SESNUMBER12' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SESNUMBER12',str(SES_Number_Var1.get()))
                                        if 'SESNU3MBER12' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SESNU3MBER12',str(SES_Number_Var2.get()))
                                        if 'AMMOUNT1SES' in paragraph.text:
                                            myammount1 = str(Ammount_Var1.get()).replace(",","")
                                            paragraph.text = paragraph.text.replace('AMMOUNT1SES',special_format(float(myammount1)))
                                        if 'AMMO12UNT1SES' in paragraph.text:
                                            myammount2 = str(Ammount_Var2.get()).replace(",","")
                                            paragraph.text = paragraph.text.replace('AMMO12UNT1SES',special_format(float(myammount2))) 
                                        if 'HHHHHHHHHHHH' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('HHHHHHHHHHHH',str(Senior_Officer.get()))
                                        if 'DDDDDDDDDDD' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('DDDDDDDDDDD',str(Senior_Designation.get()))
                                        if 'VENDORNAMEX' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('VENDORNAMEX',str(empty_list[5]))
                                        if 'SESSPEIOD' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SESSPEIOD',str(SES1_start_date))
                                        if 'TOTALXAMOUNT' in paragraph.text:
                                            #paragraph.text = paragraph.text.replace('AMMOUNT1SES',str(Ammount_Var1.get()))
                                            x1 = Ammount_Var1.get().replace(",","")
                                            x1 = float(x1)
                                            x2 = Ammount_Var2.get().replace(",","")
                                            x2 = float(x2)
                                            x3 = x1 + x2
                                            paragraph.text = paragraph.text.replace('TOTALXAMOUNT',special_format(float(x3)))
                                            #paragraph.add_run(str(x3)).font.size = Pt(10)
                                        if 'SESEPEIOD' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('SESEPEIOD',str(SES1_end_date))
                                        if 'PONUMBER' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('PONUMBER',str(PO_Number_Var.get()))
                                        if 'POSTARTDATE' in paragraph.text:
                                            paragraph.text = paragraph.text.replace('POSTARTDATE',str(PODATE))
                                        style = document.styles['Normal']
                                        font = style.font
                                        font.name = 'Calibri'
                                        font.size = Pt(10)
                                        paragraph.style = document.styles['Normal']
                            
                                    for table in document.tables:
                                        for row in table.rows:
                                            for cell in row.cells:
                                                for paragraph in cell.paragraphs:
                                                    if 'TODAYDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('TODAYDATE',str(today))
                                                    if 'HARDWARENAME' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('HARDWARENAME',str(empty_list[1]))
                                                    if 'SESTARTDATE1' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESTARTDATE1',"")
                                                        paragraph.add_run(str(Service_s_date)).font.size = Pt(7)
                                                    if 'SEENDDATE1' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SEENDDATE1',"")
                                                        paragraph.add_run(str(Service_e_date)).font.size = Pt(7)
                                                    if 'SESTARTDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESTARTDATE',str(Service_s_date))
                                                        #paragraph.add_run(str(Service_s_date)).font.size = Pt(7)
                                                    if 'SEENDDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SEENDDATE',str(Service_e_date))
                                                        #paragraph.add_run(str(Service_e_date)).font.size = Pt(7)
                                                    if 'CONTRACTNUMBER' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('CONTRACTNUMBER',str(empty_list[0]))
                                                    if 'CONTRACTSTARTDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('CONTRACTSTARTDATE',str(empty_list[2]))
                                                    if 'CONTRACTENDDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('CONTRACTENDDATE',str(empty_list[3]))
                                                    if 'VENDORCODE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('VENDORCODE',str(empty_list[4]))
                                                    if 'BILLNUMBERX' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('BILLNUMBERX',str(Bill_var.get()))
                                                    if 'BILLDATEX' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('BILLDATEX',str(Invoice_date))
                                                    if 'SESNUMBER12' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESNUMBER12',"")
                                                        paragraph.add_run(str(SES_Number_Var1.get())).font.size = Pt(9)
                                                    if 'SESNU123MBER' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESNU123MBER',"")
                                                        paragraph.add_run(str(SES_Number_Var2.get())).font.size = Pt(9)
                                                    if 'TOTALAMOUNT' in paragraph.text:
                                                        #paragraph.text = paragraph.text.replace('AMMOUNT1SES',str(Ammount_Var1.get()))
                                                        paragraph.text = paragraph.text.replace('TOTALAMOUNT',"")
                                                        x1 = Ammount_Var1.get().replace(",","")
                                                        x1 = float(x1)
                                                        x2 = Ammount_Var2.get().replace(",","")
                                                        x2 = float(x2)
                                                        x3 = x1 + x2
                                                        paragraph.add_run(special_format(float(x3))).font.size = Pt(10)
                                                    if 'HHHHHHHHHHHH' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('HHHHHHHHHHHH',str(Senior_Officer.get()))
                                                    if 'DDDDDDDDDDDD' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('DDDDDDDDDDDD',str(Senior_Designation.get()))
                                                    if 'VENDORNAMEX' in paragraph.text:
                                                        #paragraph.text = paragraph.text.replace('VENDORNAMEX',str(i[5]))
                                                        paragraph.text = paragraph.text.replace('VENDORNAMEX',"")
                                                        paragraph.add_run(str(empty_list[5])).font.size = Pt(10)
                                                    if 'SESSPEIODXY' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESSPEIODXY-SESEPEIODXY',"")
                                                        paragraph.add_run(str(SES1_start_date) + ' - ' +str(SES1_end_date)).font.size = Pt(7.5)
                                                    if 'SESSPEIOD1XY' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESSPEIOD1XY-SESEPEIOD1XY',"")
                                                        paragraph.add_run(str(SES2_start_date) + ' - ' +str(SES2_end_date)).font.size = Pt(7.5)      
                                                    if 'SESTXYARTDATE' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESTXYARTDATE-SEEXYNDDATE',"")
                                                        paragraph.add_run(str(Service_s_date) + ' - ' + str(Service_e_date)).font.size = Pt(7.5)                                        
                                                    if 'SESSPEIOD' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESSPEIOD',str(SES1_start_date))
                                                        #paragraph.text = paragraph.text.replace('SESSPEIOD',"")
                                                        #paragraph.add_run(str(SES1_start_date)).font.size = Pt(8)
                                                    if 'SESEPEIOD' in paragraph.text:
                                                        paragraph.text = paragraph.text.replace('SESEPEIOD',str(SES1_end_date))
                                                        #paragraph.text = paragraph.text.replace('SESEPEIOD',"")
                                                        #paragraph.add_run(str(SES1_end_date)).font.size = Pt(8)
                                                    if 'PONUMBER' in paragraph.text:
                                                        #paragraph.text = paragraph.text.replace('PONUMBER',str(PO_Number_Var.get()))
                                                        paragraph.text = paragraph.text.replace('PONUMBER',"")
                                                        paragraph.add_run(str(PO_Number_Var.get())).font.size = Pt(7.5)
                                                    if 'POSTARTDATE' in paragraph.text:
                                                        #paragraph.text = paragraph.text.replace('POSTARTDATE',str(PODATE))
                                                        paragraph.text = paragraph.text.replace('POSTARTDATE',"")
                                                        paragraph.add_run(str(PODATE)).font.size = Pt(8)     
                                    file = filedialog.asksaveasfile(title = "Save a bill as", filetypes = (("word file","*.docx"),("word file","*.doc")))
                                    new_file = file.name
                                    document.save(new_file)
                                    conn.close()
                                    mylist1 = str(new_file).split('/')
                                    name = mylist1.pop()
                                    mylist = str(new_file).split('/')
                                    mylist.pop()
                                    myjoinlist = '/'.join(mylist)
                                    os.chdir(str(myjoinlist))
                                    subprocess.Popen('start winword ' + str(name),shell = True)
                                    confirm1 = tkinter.messagebox.askquestion('Warning',"Do you want to close the window???", icon='warning')
                                    if confirm1 == 'yes':
                                        Bill.after(1000, lambda: Bill.destroy())
                        
        else:
            tkinter.messagebox.showinfo("Warning","Fill the values once again")
            confirm1 = tkinter.messagebox.askquestion('Warning',"Do you want to close the window???", icon='warning')
            if confirm1 == 'yes':
                Bill.after(1000, lambda: Bill.destroy())
            


    PO_Number_Var = StringVar()
    Bill_var = StringVar()
    Ammount_Var1 = StringVar()
    Ammount_Var2 = StringVar()
    SES_Number_Var1 = StringVar()
    SES_Number_Var2 = StringVar()
    Senior_Officer = StringVar()
    Senior_Designation = StringVar()
    list_var = StringVar(Bill)
    Entry_date1 = StringVar()
    Entry_service_start_date1 = StringVar()
    Entry_service_end_date1 = StringVar()
    Entry_SES1_start_date1 = StringVar()
    Entry_SES1_end_date1 = StringVar()
    Entry_SES2_start_date1 = StringVar()
    Entry_SES2_end_date1 = StringVar()
    Entry_PO_Date1 = StringVar()

                    
    conn = sqlite3.connect(db_file)
    if checkTableExists(conn, 'AMC_MANAGE_TABLE'):
        pass
    else :
        conn.execute("create table AMC_MANAGE_TABLE (CONTRACT_NO varchar(60) PRIMARY KEY, HARDWARE_NAME varchar(70), STARTING_DATE varchar(20),ENDING_DATE varchar(20), VENDOR_CODE varchar(80), VENDOR_NAME varchar(80))")
        conn.execute("INSERT INTO AMC_MANAGE_TABLE (CONTRACT_NO,HARDWARE_NAME,STARTING_DATE,ENDING_DATE,VENDOR_CODE,VENDOR_NAME)                        VALUES ( 0,null,null,null,null,null)")
    c = conn.cursor()            
    pointer = c.execute("select HARDWARE_NAME from AMC_MANAGE_TABLE")
    conn.commit()
    elist = []
    for i in pointer:
        elist.append(i[0])
    list_var.set(elist[0])
    print(elist)
    #Implementing the dropdown
    menu = OptionMenu(Bill,list_var,*elist)
    menu.grid(row = 0, column = 1, padx = 10, pady = 10)
    

    def find_details():
        mylist = list_var.get()
        if mylist =="None":
            tkinter.messagebox.showinfo("Warning","Hardware name can't be none")
        else:
            conn = sqlite3.connect(db_file)
            if checkTableExists(conn, 'AMC_MANAGE_TABLE'):
                pass
            else:
                conn.execute("create table AMC_MANAGE_TABLE (CONTRACT_NO varchar(60) PRIMARY KEY, HARDWARE_NAME varchar(70), STARTING_DATE varchar(20),ENDING_DATE varchar(20), VENDOR_CODE varchar(80), VENDOR_NAME varchar(80))")
                conn.execute("INSERT INTO AMC_MANAGE_TABLE (CONTRACT_NO,HARDWARE_NAME,STARTING_DATE,ENDING_DATE,VENDOR_CODE,VENDOR_NAME)                        VALUES ( 0,null,null,null,null,null)")
            c = conn.cursor()
            c.execute("Select *from AMC_MANAGE_TABLE where HARDWARE_NAME = ?",(mylist,))
            fetch = c.fetchone()
            empty_list = []
            if fetch == None:
                tkinter.messagebox.showinfo("Warning","Hardware name can't be none")
            else:
                for row in fetch:
                    empty_list.append(row)
                    
            contract_var = StringVar()
            contract_var.set(empty_list[0])
            Contract = Label(Bill, text = 'Contract No.')
            Contract.grid(row = 1,column = 0,padx = 10, pady = 10)
            entry_contract = Entry(Bill,textvariable = contract_var,state='disabled')
            entry_contract.grid(row = 1,column = 1,padx = 10,pady = 10)
            
            starting_var = StringVar()
            starting_var.set(empty_list[2])
            start = Label(Bill, text = 'Start date')
            start.grid(row = 1,column =2 ,padx = 10, pady = 10)
            entry_start = Entry(Bill,textvariable = starting_var,state='disabled')
            entry_start.grid(row = 1,column = 3,padx = 10,pady = 10)

            ending_var = StringVar()
            ending_var.set(empty_list[3])
            end = Label(Bill, text = 'End date')
            end.grid(row = 1,column =4 ,padx = 10, pady = 10)
            entry_end = Entry(Bill,textvariable = ending_var,state='disabled')
            entry_end.grid(row = 1,column = 5,padx = 10,pady = 10)
            
            vendor_code_var = StringVar()
            vendor_code_var.set(empty_list[4])
            v_code = Label(Bill, text = 'Vendor Code')
            v_code.grid(row = 1,column =6 ,padx = 10, pady = 10)
            entry_code = Entry(Bill,textvariable = vendor_code_var,state='disabled')
            entry_code.grid(row = 1,column = 7,padx = 10,pady = 10)
            
            vendor_name_var = StringVar()
            vendor_name_var.set(empty_list[5])
            v_name = Label(Bill, text = 'Vendor Name')
            v_name.grid(row = 1,column =8 ,padx = 10, pady = 10)
            entry_name = Entry(Bill,textvariable = vendor_name_var,state='disabled')
            entry_name.grid(row = 1,column = 9,padx = 10,pady = 10)
            

    choose_button = Button(Bill, text = 'View Details',command = find_details)
    choose_button.grid(row = 0 , column = 2,padx = 10, pady = 10)
    
    Hardware_name = Label(Bill, text = 'Select Hardware')
    Hardware_name.grid(row = 0, sticky = E, padx = 10,pady = 10)
    
    PO_No = Label(Bill, text = 'P.O. Number')
    PO_No.grid(row = 2 , sticky = E, padx = 10 , pady = 10)
    
    PO_DATE = Label(Bill, text = 'P.O. Date')
    PO_DATE.grid(row = 2, column = 3,sticky = E, padx  = 10, pady = 10)

    format1 = Label(Bill, text = '(dd/mm/yyyy)')
    format1.grid(row = 2,column = 5,sticky=E,padx = 10, pady = 10)
    

    #This is for the invoice number
    Bill_No = Label(Bill, text = 'Invoice No.')
    Bill_No.grid(row = 3, sticky = E,padx = 10, pady=10)
    
    #This is for the invoice start date
    Start_date = Label(Bill, text = 'Invoice Date')
    Start_date.grid(row = 4 , sticky = E, padx = 5 , pady = 10)

    format2 = Label(Bill, text = '(dd/mm/yyyy)')
    format2.grid(row = 4,column = 2,sticky=E)
    
    
    #This is for the service start date
    Service_Start_date = Label(Bill, text = 'Service Start Date')
    Service_Start_date.grid(row = 5 , sticky = E, padx = 10 , pady = 10)
    format2 = Label(Bill, text = 'PRESS ENTER')
    format2.grid(row = 5,column = 2,sticky=E)
    

    
    #This is for the service end date
    Service_End_date = Label(Bill, text = 'Service End Date')
    Service_End_date.grid(row = 6 , sticky = E, padx = 10 , pady = 10)
    format3 = Label(Bill, text = '(dd/mm/yyyy)')
    format3.grid(row = 6,column = 2,sticky=E)
    
   
    SES1 = Label(Bill, text = 'SES Number1')
    SES1.grid(row = 8 , sticky = E, padx = 10 , pady=(30,10))
    
    SES1_AMMOUNT = Label(Bill, text = 'SES1 AMMOUNT')
    SES1_AMMOUNT.grid(row = 8,column = 3, sticky = E, padx = 10 , pady=(30,10))

    SES1_START_DATE = Label(Bill, text = 'SES1 Start Date')
    SES1_START_DATE.grid(row = 9, sticky = E, padx = 10 , pady = 10)
    format4 = Label(Bill, text = '(dd/mm/yyyy)')
    format4.grid(row = 9,column = 2,sticky=E)
    
    
    SES1_END_DATE = Label(Bill, text = 'SES1 End Date')
    SES1_END_DATE.grid(row = 9, column = 3,sticky = E, padx = 10 , pady = 10)
    format5 = Label(Bill, text = '(dd/mm/yyyy)')
    format5.grid(row = 9,column = 5,sticky=E)
    
    SES2 = Label(Bill, text = 'SES Number2')
    SES2.grid(row = 11 , sticky = E, padx = 10 , pady = 10)
    
    SES2_AMMOUNT = Label(Bill, text = 'SES2 AMMOUNT')
    SES2_AMMOUNT.grid(row = 11,column = 3, sticky = E, padx = 10 , pady = 10)

    SES2_START_DATE = Label(Bill, text = 'SES2 Start Date')
    SES2_START_DATE.grid(row = 12, sticky = E, padx = 10 , pady = 10)
    format5 = Label(Bill, text = '(dd/mm/yyyy)')
    format5.grid(row = 12,column = 2,sticky=E)
    
    
    SES2_END_DATE = Label(Bill, text = 'SES2 End Date')
    SES2_END_DATE.grid(row = 12, column = 3,sticky = E, padx = 10 , pady = 10)
    format5 = Label(Bill, text = '(dd/mm/yyyy)')
    format5.grid(row = 12,column = 5,sticky=E)
    
    Senior_officer = Label(Bill, text = "Officer's Name")
    Senior_officer.grid(row = 14 , sticky = E, padx = 10 , pady = (30,10))

    Senior_officer_Designation = Label(Bill, text = "Officer's Designation")
    Senior_officer_Designation.grid(row = 15 , sticky = E, padx = 10 , pady = 10)
    
    conn = sqlite3.connect(db_file)
    if checkTableExists(conn, 'bill'):
        pass
    else :
        conn.execute('create table bill (id INTEGER PRIMARY KEY autoincrement, PO_Number varchar(100), PO_DATE varchar(20), Invoice_Number varchar(100), Invoice_date varchar(20), Service_start_date varchar(20), Service_end_date varchar(20),SES_number1 varchar(60), SES1_Ammount bigint, SES1_start_date varchar(20),SES1_end_date varchar(20),SES_number2 varchar(60), SES2_Ammount bigint, SES2_start_date varchar(20),SES2_end_date varchar(20),Senior_officer varchar(100), Designation varchar(100))')
        conn.execute("INSERT INTO bill (PO_Number,PO_DATE,Invoice_Number, Invoice_date, Service_start_date, Service_end_date ,SES_number1 , SES1_Ammount, SES1_start_date ,SES1_end_date ,SES_number2, SES2_Ammount , SES2_start_date,SES2_end_date,Senior_officer, Designation)                        VALUES (null,null,null,null,null,null,null,null,null,null,null,null,null,null,null,null)");
        conn.commit()
                
        
    count = conn.execute('SELECT * FROM bill ORDER BY id DESC LIMIT 1')
    for i in count: 
        PO_Number_Var.set(i[1])
        Entry_PO_Number = Entry(Bill,textvariable = PO_Number_Var)
        Entry_PO_Number.grid(row = 2, column =1,padx = 10, pady=10)
        
        Entry_PO_Date1.set(i[2])
        Entry_PO_Date = Entry(Bill, textvariable = Entry_PO_Date1)
        Entry_PO_Date.grid(row = 2,column =4,padx =  10, pady =10)
        #Entry_PO_Date = DateEntry(Bill, width=18,height = 10, background='darkblue',foreground='white', borderwidth=2)
        #Entry_PO_Date.gr id(row = 1,column =4, padx = 10, pady = 10)
        
        Bill_var.set(i[3])
        Entry_Bill = Entry(Bill,textvariable = Bill_var)  #For the invoice no.
        Entry_Bill.grid(row = 3, column =1,padx = 10, pady=10)

        Entry_date1.set(i[4])
        Entry_date = Entry(Bill, textvariable = Entry_date1)
        #Entry_date = DateEntry(Bill, width=18,height = 10, background='darkblue',foreground='white', borderwidth=2) #For the invoice date
        Entry_date.grid(row = 4,column = 1,padx=10, pady=10)

        #Entry_service_start_date1.set(i[5])
        def change_date(event):
            date = Entry_service_start_date1.get()
            date_obj = datetime.datetime.strptime(date,'%d/%m/%Y').date()
            date_obj = date_obj + relativedelta(months = 3)
            date_obj = date_obj - relativedelta(days = 1)
            mydate = str(date_obj).split('-')
            mydate.reverse()
            mydate = '/'.join(mydate)
            Entry_service_end_date1.set(mydate)
            Entry_SES1_start_date1.set(date)
            Entry_SES1_end_date1.set(mydate)

        Entry_service_start_date1.set('dd/mm/yyyy')    
        Entry_service_start_date = Entry(Bill, textvariable = Entry_service_start_date1)
        Entry_service_start_date.grid(row = 5,column = 1,padx=10, pady=10)
        Entry_service_start_date.bind('<Return>',change_date)

        
        Entry_service_end_date = Entry(Bill, textvariable = Entry_service_end_date1)
        Entry_service_end_date.grid(row = 6,column = 1,padx=10, pady=10)
        
        SES_Number_Var1.set(i[7])
        Entry_SES_Number1 = Entry(Bill,textvariable = SES_Number_Var1)
        Entry_SES_Number1.grid(row = 8, column =1,padx = 10, pady=(30,10))
        
        Ammount_Var1.set(i[8])
        Entry_Ammount1 = Entry(Bill,textvariable = Ammount_Var1)
        Entry_Ammount1.grid(row = 8, column =4,padx = 10, pady=(30,10))
        
        Entry_SES1_start_date = Entry(Bill, textvariable = Entry_SES1_start_date1)
        Entry_SES1_start_date.grid(row = 9,column = 1,padx=10, pady=10)
        
        Entry_SES1_end_date = Entry(Bill, textvariable = Entry_SES1_end_date1)
        Entry_SES1_end_date.grid(row = 9,column = 4,padx=10, pady=10)
        
        SES_Number_Var2.set(i[11])
        Entry_SES_Number2 = Entry(Bill,textvariable = SES_Number_Var2)
        Entry_SES_Number2.grid(row = 11, column =1,padx = 10, pady= 10)
        
        Ammount_Var2.set(i[12])
        Entry_Ammount2 = Entry(Bill,textvariable = Ammount_Var2)
        Entry_Ammount2.grid(row = 11, column =4,padx = 10, pady=10)
        
        Entry_SES2_start_date = Entry(Bill, textvariable = Entry_SES2_start_date1)
        Entry_SES2_start_date.grid(row = 12,column = 1,padx=10, pady=10)
        
        Entry_SES2_end_date = Entry(Bill, textvariable = Entry_SES2_end_date1)
        Entry_SES2_end_date.grid(row = 12,column = 4,padx=10, pady=10)
        
    
        Senior_Officer.set(i[15])
        Entry_Officer = Entry(Bill,textvariable = Senior_Officer)
        Entry_Officer.grid(row = 14, column =1,padx = 10, pady=(30,10))
        
        Senior_Designation.set(i[16])
        Entry_Officer = Entry(Bill,textvariable = Senior_Designation)
        Entry_Officer.grid(row = 15, column =1,padx = 10, pady=10)
        
    conn.close() 
    Button(Bill, text = 'Submit',command = Printing_Bills).grid(columnspan = 5,padx = 10, pady=10)
    Bill.mainloop()


    
#This is the function for the AMC whenever the user clicks over the NEW Entry for AMC this will be displayed
def AMC_Function():

    #Checking for that field should not be null
    def check_NULL(Contract,Hardware, Starting,Ending, Vendor_code, Vendor_name):
        if Ending!="" and Starting!="":
            date_list1 = Starting.split('/')
            date_list2 = Ending.split('/')
        if Contract == "" or Hardware == "" or Starting == "" or Vendor_code == "" or Vendor_name == "" or Ending == "":
            tkinter.messagebox.showinfo("Field","Fields can't be empty")
            return False
        else:
            tkinter.messagebox.showinfo('Correct','All fields are ok')
            return True
        
        
    #Sending the data to the database when the data is received from the AMC form
    def send_to_database():
        confirm = tkinter.messagebox.askquestion('Warning',"Are You Sure???", icon='warning')
        if confirm=='yes':
            Contract = Contract_var.get()
            Hardware = Hardware_Var.get()
            Starting = cal.get_date()
            Ending = cal1.get_date()
            Vendor_code = Vendor_code_Var.get()
            Vendor_name = Vendor_name_Var.get()
            lst = str(Starting).split('-')
            lst.reverse()
            Starting = '/'.join(lst)
            lst = str(Ending).split('-')
            lst.reverse()
            Ending = '/'.join(lst)

            
            conn1 = sqlite3.connect(db_file)
            if checkTableExists(conn1, 'AMC_MANAGE_TABLE'):    
                pass
            else :
                conn.execute("create table AMC_MANAGE_TABLE (CONTRACT_NO varchar(60) PRIMARY KEY, HARDWARE_NAME varchar(70), STARTING_DATE varchar(20),ENDING_DATE varchar(20), VENDOR_CODE varchar(80), VENDOR_NAME varchar(80))")
                conn.execute("INSERT INTO AMC_MANAGE_TABLE (CONTRACT_NO,HARDWARE_NAME,STARTING_DATE,ENDING_DATE,VENDOR_CODE,VENDOR_NAME)                        VALUES ( 0,null,null,null,null,null)")

            check_hardware = conn1.execute("select HARDWARE_NAME from AMC_MANAGE_TABLE")
            checking_counter = 0
            empty1_list = []
            for cursor in check_hardware:
                empty1_list.append(cursor[0])
            for check in empty1_list:
                if check ==Hardware:
                    checking_counter = 1
            conn1.close
            if checking_counter ==0:
                if (check_NULL(Contract,Hardware, Starting, Ending, Vendor_code, Vendor_name)):
                    #Send the values to the database
                    
                    conn = sqlite3.connect(db_file)
                    if checkTableExists(conn, 'AMC_MANAGE_TABLE'):
                        pass
                    else :
                        conn.execute("create table AMC_MANAGE_TABLE (CONTRACT_NO varchar(60) PRIMARY KEY, HARDWARE_NAME varchar(70), STARTING_DATE varchar(20),ENDING_DATE varchar(20), VENDOR_CODE varchar(80), VENDOR_NAME varchar(80))")
                    conn.execute("INSERT INTO AMC_MANAGE_TABLE (CONTRACT_NO,HARDWARE_NAME,STARTING_DATE,ENDING_DATE,VENDOR_CODE,VENDOR_NAME)                        VALUES ( ?,?,?,?,?,?)",(Contract, Hardware, Starting, Ending,Vendor_code, Vendor_name))
                    conn.commit()
                    pointer2 = conn.execute("select HARDWARE_NAME from AMC_MANAGE_TABLE")
                    #list_var2 = StringVar()
                    emp_list1 = []
                    for i in pointer2:
                        emp_list1.append(i[0])
                    list_var1.set(emp_list1[0])
                    menu1 = OptionMenu(root,list_var1,*emp_list1)
                    menu1.grid(padx = 10, pady = 10)
                    menu1.place(x = 880, y = 300)
                    conn.close()
                    tkinter.messagebox.showinfo('Entered','Values entered successfully')
                    #To clear the field
                    Entry_Contract.delete(0,END)
                    Entry_Hardware.delete(0,tkinter.END)
                    Entry_VCode.delete(0,tkinter.END)
                    Entry_Vname.delete(0,tkinter.END)
                    confirm1 = tkinter.messagebox.askquestion('Warning',"Do you want to close the window???", icon='warning')
                    if confirm1 == 'yes':
                        AMC.after(1000, lambda: AMC.destroy())
                
                else :
                    tkinter.messagebox.showinfo('Fill Values','Please fill the correct values')
                    AMC.after(1000, lambda: AMC.destroy())
            else:
                tkinter.messagebox.showinfo('Name Error','Hardware Name already exist, please choose different Hardware Name')
                AMC.after(1000, lambda: AMC.destroy())
        else:
            AMC.after(1000, lambda: AMC.destroy())
     
    #Window of New entry for AMC
    AMC = Toplevel()
    AMC.geometry("600x350+300+300")
    
    #To save the enterd values of the fields
    Contract_var = StringVar()
    Hardware_Var = StringVar()
    Vendor_code_Var = StringVar()
    Vendor_name_Var = StringVar()
    
    Contract_No = Label(AMC, text = 'Contract No.')
    Contract_No.grid(row = 0 , sticky = E,padx = 10, pady=10)
    
    Hardware_Name = Label(AMC, text = 'Hardware Name')
    Hardware_Name.grid(row = 1, sticky = E,padx = 10, pady=10)
    
    Starting_date = Label(AMC, text = 'Contract Start date')
    Starting_date.grid(row = 2, sticky = E,padx = 10, pady=10)
    
    End_date = Label(AMC, text = 'Contract End date')
    End_date.grid(row = 3, sticky = E,padx = 10, pady=10)
    
    Vendor_Code = Label(AMC, text = 'Vendor Code')
    Vendor_Code.grid(row = 4, sticky = E,padx = 10, pady=10)
    
    Vendor_Name = Label(AMC, text = 'Vendor_Name')
    Vendor_Name.grid(row = 5, sticky = E,padx = 10, pady=10)
                                                        
    
    #To take the entries of the variables
    Entry_Contract = Entry(AMC,textvariable = Contract_var)
    Entry_Contract.grid(row = 0, column =1,padx = 10, pady=10)
    
    Entry_Hardware = Entry(AMC,textvariable = Hardware_Var)
    Entry_Hardware.grid(row = 1, column =1,padx = 10, pady=10)
    
    cal = DateEntry(AMC, width=18,height = 10, background='darkblue',foreground='white', borderwidth=2)
    cal.grid(row = 2,column = 1,padx=10, pady=10)
    
    cal1 = DateEntry(AMC, width=18,height = 10, background='darkblue',foreground='white', borderwidth=2)
    cal1.grid(row = 3,column = 1,padx=10, pady=10)
        
    Entry_VCode = Entry(AMC,textvariable = Vendor_code_Var)
    Entry_VCode.grid(row = 4, column =1,padx = 10, pady=10)
    
    Entry_Vname = Entry(AMC,textvariable = Vendor_name_Var)
    Entry_Vname.grid(row = 5, column =1,padx = 10, pady=10)
                                                        
    Button(AMC, text = 'Submit',command = send_to_database).grid(padx = 10, pady=10)
    AMC.mainloop()
    
root= Tk()
root.state('zoomed')
root.title("AMC Generator")
root.configure(background='Orange')
temp = Image.open('index.png')
temp = temp.resize((150, 150), Image.ANTIALIAS)
temp.save("xyz.png", "png")
bg_image1 = PhotoImage(file ="xyz.png")
img1 = Label(root,image = bg_image1)
img1.pack()
img1.place(x =1200,y = 20)

temp1 = Image.open('ashok.png')
temp1 = temp1.resize((150, 150), Image.ANTIALIAS)
temp1.save("pqr.png", "png")
bg_image2 = PhotoImage(file ="pqr.png")
img2 = Label(root,image = bg_image2)
img2.pack()
img2.place(x =10,y = 20)

#temp2 = Image.open('oil.png')
#temp2 = temp2.resize((200,200 ), Image.ANTIALIAS)
#temp2.save("pq.png", "png")
bg_image3 = PhotoImage(file ="oil.png")
img3 = Label(root,image = bg_image3)
img3.pack()
img3.place(x =330,y = 20)

AMC_Button=Button(root,text="New Entry for AMC Contract / Agreement",command=AMC_Function,padx = 10, pady=10)
AMC_Button.pack()
AMC_Button.place(x = 370, y = 300)
Bill_Genearte=Button(root,text="Generate Invoice/ Bill",command = Bill_generator_function, padx = 10, pady=10)
Bill_Genearte.pack()
Bill_Genearte.place(x = 420, y = 400)

list_var1 = StringVar(root)
conn = sqlite3.connect(db_file)
if checkTableExists(conn, 'AMC_MANAGE_TABLE'):
    pass
else :
    conn.execute("create table AMC_MANAGE_TABLE (CONTRACT_NO varchar(60) PRIMARY KEY, HARDWARE_NAME varchar(70), STARTING_DATE varchar(20),ENDING_DATE varchar(20), VENDOR_CODE varchar(80), VENDOR_NAME varchar(80))")
    conn.execute("INSERT INTO AMC_MANAGE_TABLE (CONTRACT_NO,HARDWARE_NAME,STARTING_DATE,ENDING_DATE,VENDOR_CODE,VENDOR_NAME)                        VALUES ( 0,null,null,null,null,null)")
                

pointer1 = conn.execute("select HARDWARE_NAME from AMC_MANAGE_TABLE")
emp_list = []
for i in pointer1:
    emp_list.append(i[0])
    
list_var1.set(emp_list[0])

menu1 = OptionMenu(root,list_var1,*emp_list)
menu1.grid(padx = 10, pady = 10)
menu1.place(x = 880, y = 300)
conn.commit()
conn.close()

def delete_entry():
    confirm = tkinter.messagebox.askquestion('Warning',"Are You Sure you want to delete this entry???", icon='warning')
    if confirm =='yes':
        temp = list_var1.get()
        if temp=='None':
            tkinter.messagebox.showinfo("Empty","Hardware name is not selected")
        else:
            conn = sqlite3.connect(db_file)
            c = conn.cursor()
            c.execute('delete from AMC_MANAGE_TABLE where HARDWARE_NAME = (?)',(temp,))
            conn.commit()
            tkinter.messagebox.showinfo("Success","Deleted Successfully")
            pointer1 = conn.execute("select HARDWARE_NAME from AMC_MANAGE_TABLE")
            emp_list = []
            for i in pointer1:
                emp_list.append(i[0])
            list_var1.set(emp_list[0])
            menu1 = OptionMenu(root,list_var1,*emp_list)
            menu1.grid(padx = 10, pady = 10)
            menu1.place(x = 880, y = 300)
    else:
        tkinter.messagebox.showinfo("Alert","Ok Not deleted")
        return 0

del_button = Button(root,text = 'Delete',command = delete_entry)
del_button.grid(padx = 20, pady = 20)
del_button.config(width = 7)
del_button.place(x = 880, y = 360)
root.mainloop()
